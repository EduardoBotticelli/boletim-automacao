"""
Gera os arquivos Word do boletim a partir de boletim.json.
- boletim.docx (geral, com TODOS os itens)
- boletim_<slug>.docx (um por area, filtrado pelo campo 'boletins' de cada item)

Uso: roda APOS gerar_boletim.py no mesmo workflow.
Layout provisorio inspirado no HTML (aprovado pela Alice) - sera adaptado
ao template oficial da equipe de Marketing quando estiver disponivel.
"""

import os
import json
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
INPUT_PATH = os.path.join(BASE_DIR, "output", "boletim.json")
OUTPUT_DIR = os.path.join(BASE_DIR, "output")

if not os.path.exists(INPUT_PATH):
    raise SystemExit("ERRO: boletim.json nao encontrado em " + INPUT_PATH)

with open(INPUT_PATH, "r", encoding="utf-8") as f:
    boletim = json.load(f)

data_exec_iso = boletim.get("data_execucao", "")
todos_itens = boletim.get("itens", [])
sem_publicacao = boletim.get("fontes_sem_publicacao_hoje", [])
sem_resultado = boletim.get("fontes_sem_resultado", [])
com_erro = boletim.get("fontes_com_erro_tecnico", [])
janela = boletim.get("janela_aplicada", {})
config = boletim.get("boletins_config", {})
fontes_email_pendentes = config.get("fontes_email_pendentes", {})

# Nomes bonitos para exibicao
NOMES_BONITOS = {
    "trabalhista": "Trabalhista",
    "tributario": "Tributario",
    "empresarial": "Empresarial, M&A e Mercado de Capitais",
    "regulatorio": "Regulatorio",
    "imobiliario": "Imobiliario e Infraestrutura",
    "ambiental": "Ambiental e ESG",
    "propriedade-intelectual": "Propriedade Intelectual, Tecnologia e Privacidade",
    "contencioso": "Contencioso",
}

meses_pt = ["janeiro", "fevereiro", "marco", "abril", "maio", "junho",
            "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
try:
    d = datetime.date.fromisoformat(data_exec_iso)
    data_extenso = str(d.day) + " de " + meses_pt[d.month - 1] + " de " + str(d.year)
    data_curta = d.strftime("%d/%m/%Y")
except Exception:
    data_extenso = data_exec_iso
    data_curta = data_exec_iso


# ---------------------- HELPERS DE ESTILO ----------------------

VERDE_ESCURO = RGBColor(0x1A, 0x4D, 0x2E)
VERDE_CLARO = RGBColor(0x2D, 0x86, 0x59)
VERDE_MUITO_CLARO = RGBColor(0xA8, 0xD8, 0xB9)
VERMELHO_ALTA = RGBColor(0xC0, 0x39, 0x2B)
LARANJA_MEDIA = RGBColor(0xD6, 0x89, 0x10)
CINZA_BAIXA = RGBColor(0x7F, 0x8C, 0x8D)
CINZA_TEXTO = RGBColor(0x4A, 0x4A, 0x4A)
CINZA_MEDIO = RGBColor(0x88, 0x88, 0x88)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
PRETO = RGBColor(0x1A, 0x1A, 0x1A)
FUNDO_AVISO = "FEF5E7"
FUNDO_VERDE = "1A4D2E"
FUNDO_VERDE_CLARO = "F0F7F2"


def set_cell_shading(cell, hex_color):
    """Aplica cor de fundo a uma celula."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_hyperlink(paragraph, url, text, color=None, bold=False):
    """Adiciona um hyperlink clicavel a um paragrafo."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    if bold:
        b = OxmlElement("w:b")
        r_pr.append(b)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def rgb_to_hex(rgb):
    return "{:02X}{:02X}{:02X}".format(rgb[0], rgb[1], rgb[2])


# ---------------------- CONSTRUCAO DO DOCUMENTO ----------------------

def gerar_word(itens, titulo_boletim, subtitulo_extra, aviso_email_pendente, output_path):
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Estilo padrao
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---------- CAPA ----------
    capa_table = doc.add_table(rows=1, cols=1)
    capa_table.autofit = False
    capa_table.columns[0].width = Cm(17)
    capa_cell = capa_table.rows[0].cells[0]
    capa_cell.width = Cm(17)
    set_cell_shading(capa_cell, FUNDO_VERDE)

    # Titulo grande "Boletim"
    p_titulo = capa_cell.paragraphs[0]
    p_titulo.paragraph_format.space_before = Pt(30)
    p_titulo.paragraph_format.space_after = Pt(6)
    run_titulo = p_titulo.add_run("Boletim")
    run_titulo.font.size = Pt(44)
    run_titulo.font.color.rgb = BRANCO
    run_titulo.font.bold = False

    # Subtitulo (categoria/area)
    p_sub = capa_cell.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(6)
    run_sub = p_sub.add_run(subtitulo_extra.upper())
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = VERDE_MUITO_CLARO
    run_sub.font.bold = True

    # Data
    p_data = capa_cell.add_paragraph()
    p_data.paragraph_format.space_after = Pt(20)
    run_data = p_data.add_run(data_extenso)
    run_data.font.size = Pt(14)
    run_data.font.color.rgb = BRANCO

    doc.add_paragraph()  # espacador

    # ---------- ESTATISTICAS ----------
    n_total = len(itens)
    n_alta = sum(1 for i in itens if i.get("relevancia") == "Alta")
    n_media = sum(1 for i in itens if i.get("relevancia") in ["Media", "Média"])
    n_baixa = sum(1 for i in itens if i.get("relevancia") == "Baixa")

    stats_table = doc.add_table(rows=2, cols=4)
    stats_table.autofit = False
    for col in stats_table.columns:
        col.width = Cm(4)

    labels = ["Total", "Alta", "Media", "Baixa"]
    valores = [n_total, n_alta, n_media, n_baixa]
    cores = [VERDE_ESCURO, VERMELHO_ALTA, LARANJA_MEDIA, CINZA_BAIXA]

    for i in range(4):
        cell_val = stats_table.rows[0].cells[i]
        cell_val.width = Cm(4)
        cell_val.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell_val.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(valores[i]))
        run.font.size = Pt(28)
        run.font.color.rgb = cores[i]
        run.font.bold = False

        cell_lbl = stats_table.rows[1].cells[i]
        cell_lbl.width = Cm(4)
        p2 = cell_lbl.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(labels[i].upper())
        run2.font.size = Pt(9)
        run2.font.color.rgb = CINZA_MEDIO
        run2.font.bold = True

    doc.add_paragraph()

    # ---------- AVISO DE FONTES POR EMAIL PENDENTES ----------
    if aviso_email_pendente:
        aviso_table = doc.add_table(rows=1, cols=1)
        aviso_table.autofit = False
        aviso_table.columns[0].width = Cm(17)
        aviso_cell = aviso_table.rows[0].cells[0]
        aviso_cell.width = Cm(17)
        set_cell_shading(aviso_cell, FUNDO_AVISO)

        p_aviso = aviso_cell.paragraphs[0]
        p_aviso.paragraph_format.space_before = Pt(8)
        p_aviso.paragraph_format.space_after = Pt(8)
        run_a1 = p_aviso.add_run("Nota: ")
        run_a1.font.bold = True
        run_a1.font.size = Pt(10)
        run_a1.font.color.rgb = LARANJA_MEDIA

        run_a2 = p_aviso.add_run(
            "Este boletim tera as seguintes fontes adicionais quando a integracao por e-mail estiver ativa: "
            + ", ".join(aviso_email_pendente) + "."
        )
        run_a2.font.size = Pt(10)
        run_a2.font.color.rgb = CINZA_TEXTO

        doc.add_paragraph()

    # ---------- BOLETIM VAZIO ----------
    if not itens:
        p_vazio = doc.add_paragraph()
        p_vazio.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_v = p_vazio.add_run("Nao ha itens para este boletim na janela temporal atual.")
        run_v.font.size = Pt(12)
        run_v.font.color.rgb = CINZA_MEDIO

        p_vazio2 = doc.add_paragraph()
        p_vazio2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_v2 = p_vazio2.add_run("Isso pode ocorrer em fins de semana, feriados ou dias sem publicacoes relevantes.")
        run_v2.font.size = Pt(10)
        run_v2.font.color.rgb = CINZA_MEDIO
        run_v2.italic = True

        _rodape(doc, janela)
        doc.save(output_path)
        return

    # ---------- DESTAQUES DO DIA (top 5 Alta) ----------
    destaques = [i for i in itens if i.get("relevancia") == "Alta"][:5]
    if destaques:
        _cabecalho_secao(doc, "Destaques do dia", "Selecionados por alta relevancia")
        for item in destaques:
            _card_destaque(doc, item)
        doc.add_paragraph()

    # ---------- AGRUPAMENTO POR CATEGORIA ----------
    categorias = {}
    ordem_categorias = []
    for item in itens:
        cat = item.get("categoria", "Outros")
        if cat not in categorias:
            categorias[cat] = []
            ordem_categorias.append(cat)
        categorias[cat].append(item)

    for cat in ordem_categorias:
        _cabecalho_secao(doc, cat, "")

        # Agrupar por fonte dentro da categoria
        fontes_na_cat = {}
        ordem_fontes = []
        for item in categorias[cat]:
            fonte = item.get("fonte", "Sem fonte")
            if fonte not in fontes_na_cat:
                fontes_na_cat[fonte] = []
                ordem_fontes.append(fonte)
            fontes_na_cat[fonte].append(item)

        for fonte in ordem_fontes:
            # Titulo da fonte
            p_fonte = doc.add_paragraph()
            p_fonte.paragraph_format.space_before = Pt(12)
            p_fonte.paragraph_format.space_after = Pt(4)
            run_fonte = p_fonte.add_run(fonte.upper())
            run_fonte.font.size = Pt(10)
            run_fonte.font.bold = True
            run_fonte.font.color.rgb = VERDE_ESCURO

            # Linha abaixo do nome da fonte
            _linha_horizontal(doc, VERDE_ESCURO)

            # Itens desta fonte
            for item in fontes_na_cat[fonte]:
                _card_item(doc, item)

        doc.add_paragraph()

    # ---------- TRANSPARENCIA ----------
    _cabecalho_secao(doc, "Transparencia da coleta", "", cor_fundo=None)

    if sem_publicacao:
        p_titulo_ = doc.add_paragraph()
        run = p_titulo_.add_run("Fontes sem publicacao na janela:")
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = VERDE_ESCURO
        for f in sem_publicacao:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(f.get("fonte", "") + " - " + f.get("motivo", ""))
            r.font.size = Pt(9)
            r.font.color.rgb = CINZA_TEXTO

    if sem_resultado:
        p_titulo_ = doc.add_paragraph()
        run = p_titulo_.add_run("Fontes sem resultado relevante:")
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = LARANJA_MEDIA
        for f in sem_resultado:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(f.get("fonte", "") + " - " + f.get("motivo", ""))
            r.font.size = Pt(9)
            r.font.color.rgb = CINZA_TEXTO

    if com_erro:
        p_titulo_ = doc.add_paragraph()
        run = p_titulo_.add_run("Fontes com erro tecnico:")
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = VERMELHO_ALTA
        for f in com_erro:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(f.get("fonte", "") + " - " + f.get("motivo", ""))
            r.font.size = Pt(9)
            r.font.color.rgb = CINZA_TEXTO

    _rodape(doc, janela)
    doc.save(output_path)


# ---------------------- COMPONENTES AUXILIARES ----------------------

def _linha_horizontal(doc, color_rgb):
    """Adiciona uma linha horizontal fina."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), rgb_to_hex((color_rgb[0], color_rgb[1], color_rgb[2])))
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _cabecalho_secao(doc, titulo, subtitulo, cor_fundo=FUNDO_VERDE):
    """Cabecalho verde para uma secao (categoria, destaques, transparencia)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(17)
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(17)
    if cor_fundo:
        set_cell_shading(cell, cor_fundo)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(titulo)
    r.font.size = Pt(16)
    r.font.color.rgb = BRANCO if cor_fundo else VERDE_ESCURO
    r.font.bold = False

    if subtitulo:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(6)
        r2 = p2.add_run(subtitulo)
        r2.font.size = Pt(9)
        r2.font.color.rgb = VERDE_MUITO_CLARO if cor_fundo else CINZA_MEDIO


def _card_destaque(doc, item):
    """Card compacto para os destaques do dia."""
    fonte = item.get("fonte", "")
    titulo = item.get("titulo", "")
    resumo = item.get("resumo", "")
    url = item.get("url", "")

    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(17)
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(17)
    set_cell_shading(cell, FUNDO_VERDE_CLARO)

    p_fonte = cell.paragraphs[0]
    p_fonte.paragraph_format.space_before = Pt(6)
    p_fonte.paragraph_format.space_after = Pt(2)
    r_fonte = p_fonte.add_run(fonte.upper())
    r_fonte.font.size = Pt(8)
    r_fonte.font.bold = True
    r_fonte.font.color.rgb = VERDE_ESCURO

    p_titulo = cell.add_paragraph()
    p_titulo.paragraph_format.space_after = Pt(2)
    if url and url != "#":
        add_hyperlink(p_titulo, url, titulo, color=rgb_to_hex((PRETO[0], PRETO[1], PRETO[2])), bold=True)
        for run in p_titulo.runs:
            run.font.size = Pt(11)
    else:
        r_titulo = p_titulo.add_run(titulo)
        r_titulo.font.size = Pt(11)
        r_titulo.font.bold = True
        r_titulo.font.color.rgb = PRETO

    p_resumo = cell.add_paragraph()
    p_resumo.paragraph_format.space_after = Pt(6)
    r_resumo = p_resumo.add_run(resumo)
    r_resumo.font.size = Pt(9)
    r_resumo.font.color.rgb = CINZA_TEXTO


def _card_item(doc, item):
    """Card completo para um item de notícia."""
    titulo = item.get("titulo", "")
    resumo = item.get("resumo", "")
    url = item.get("url", "")
    data_pub = item.get("data_publicacao", "") or "data nao identificada"
    motivo = item.get("motivo_relevancia", "")
    rel = item.get("relevancia", "")

    # Titulo + badge de relevancia em linha
    p_top = doc.add_paragraph()
    p_top.paragraph_format.space_before = Pt(8)
    p_top.paragraph_format.space_after = Pt(2)
    r_titulo = p_top.add_run(titulo + "   ")
    r_titulo.font.size = Pt(11)
    r_titulo.font.bold = True
    r_titulo.font.color.rgb = PRETO

    if rel:
        r_badge = p_top.add_run("[" + rel.upper() + "]")
        r_badge.font.size = Pt(8)
        r_badge.font.bold = True
        rel_norm = rel.replace("Média", "Media")
        if rel_norm == "Alta":
            r_badge.font.color.rgb = VERMELHO_ALTA
        elif rel_norm == "Media":
            r_badge.font.color.rgb = LARANJA_MEDIA
        else:
            r_badge.font.color.rgb = CINZA_BAIXA

    # Resumo
    p_resumo = doc.add_paragraph()
    p_resumo.paragraph_format.space_after = Pt(2)
    r = p_resumo.add_run(resumo)
    r.font.size = Pt(10)
    r.font.color.rgb = CINZA_TEXTO

    # Rodape do item: data + link
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(2)
    r_data_lbl = p_meta.add_run("Publicado: ")
    r_data_lbl.font.size = Pt(9)
    r_data_lbl.font.bold = True
    r_data_lbl.font.color.rgb = CINZA_MEDIO
    r_data = p_meta.add_run(data_pub + "   ")
    r_data.font.size = Pt(9)
    r_data.font.color.rgb = CINZA_MEDIO

    if url and url != "#":
        add_hyperlink(p_meta, url, "Acessar materia", color=rgb_to_hex((VERDE_ESCURO[0], VERDE_ESCURO[1], VERDE_ESCURO[2])), bold=True)
        for run in p_meta.runs[-1:]:
            run.font.size = Pt(9)

    # Motivo de relevancia
    if motivo:
        p_motivo = doc.add_paragraph()
        p_motivo.paragraph_format.space_after = Pt(6)
        r_m = p_motivo.add_run(motivo)
        r_m.font.size = Pt(8)
        r_m.italic = True
        r_m.font.color.rgb = CINZA_MEDIO


def _rodape(doc, janela):
    """Rodape com informacoes tecnicas."""
    doc.add_paragraph()
    _linha_horizontal(doc, CINZA_MEDIO)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Rascunho gerado automaticamente para validacao interna.")
    r.font.size = Pt(8)
    r.font.color.rgb = CINZA_MEDIO
    r.italic = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Curadoria por IA + revisao humana recomendada antes da distribuicao.")
    r2.font.size = Pt(8)
    r2.font.color.rgb = CINZA_MEDIO
    r2.italic = True

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Janela analisada: " + janela.get("inicio", "") + " ate " + janela.get("fim", ""))
    r3.font.size = Pt(8)
    r3.font.color.rgb = CINZA_MEDIO


# ---------------------- EXECUCAO ----------------------

os.makedirs(OUTPUT_DIR, exist_ok=True)

# Boletim geral
path_geral = os.path.join(OUTPUT_DIR, "boletim.docx")
gerar_word(
    itens=todos_itens,
    titulo_boletim="Boletim Juridico",
    subtitulo_extra="Curadoria de todas as fontes",
    aviso_email_pendente=[],
    output_path=path_geral,
)
print("Word geral: " + path_geral + " (" + str(len(todos_itens)) + " itens)")

# 8 boletins por area
slugs = config.get("boletins_disponiveis", [])
for slug in slugs:
    itens_do_boletim = [i for i in todos_itens if slug in i.get("boletins", [])]
    aviso = fontes_email_pendentes.get(slug, [])
    nome_bonito = NOMES_BONITOS.get(slug, slug.title())

    path_area = os.path.join(OUTPUT_DIR, "boletim_" + slug + ".docx")
    gerar_word(
        itens=itens_do_boletim,
        titulo_boletim="Boletim - " + nome_bonito,
        subtitulo_extra=nome_bonito,
        aviso_email_pendente=aviso,
        output_path=path_area,
    )
    print("Word " + slug + ": " + path_area + " (" + str(len(itens_do_boletim)) + " itens)")

print("")
print("Concluido - " + str(1 + len(slugs)) + " arquivos .docx gerados")
